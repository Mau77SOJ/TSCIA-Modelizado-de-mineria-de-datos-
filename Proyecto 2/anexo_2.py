from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Crear documento Word 
doc = Document()

# Título 
title = doc.add_heading("Ejercicio Práctico – Modelizado de Minería de Datos", 0) 
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introducción 
doc.add_paragraph("Este ejercicio está diseñado para aplicar los contenidos clave de la asignatura 'Modelizado de Minería de Datos', incluyendo preprocesamiento, visualización, modelado predictivo, y toma de decisiones basada en datos.")

# Secciones del ejercicio
sections = {
    "Objetivo": [
        "Predecir si un cliente realizará una recompra después de recibir una promoción, utilizando técnicas de modelado y visualización de datos."
    ],
    "Dataset Sugerido": [
        "Usar el archivo 'Mini_Proyecto_Clientes_Promociones.xlsx' con los siguientes campos:",
        "- Cliente_ID, Género, Edad, Recibió_Promo, Monto_Promocion, Recompra, Total_Compras, Ingreso_Mensual"
    ], 
    "Pasos para realizar el ejercicio": [
        "A continuación se detallan los pasos a seguir:"
    ],
    "1. Comprensión del Problema": [
        "Contexto: Un área de marketing quiere saber si los clientes que recibieron promociones volverán a comprar.",
        "Preguntas guía:",
        "- ¿Recibir una promoción influye en la recompra?",
        "- ¿Importa el monto?",
        "- ¿Influyen edad o ingreso?"
    ],
    "2. Carga y Exploración del Dataset": [ 
        "- Cargar los datos en Python o Excel.", 
        "- Verificar datos faltantes, valores extremos, etc.", 
        "- Analizar la distribución de las variables." 
    ], 
    "3. Transformación y Codificación": [ 
        "- Convertir variables categóricas a numéricas.", 
        "- Crear nuevas variables si es necesario." 
    ], 
    "4. Visualización de Relaciones Clave": [ 
        "- Recompra vs. Monto de promoción.", 
        "- Recompra vs. Ingreso mensual.", 
        "- Distribución por género y edad." 
    ], 
    "5. Modelado Predictivo – Clasificación": [
        "- Modelo sugerido: Árbol de Decisión.",
        "- Entrenar, predecir y evaluar el modelo." 
    ], 
    "6. Toma de Decisiones Basada en Visualización": [ 
        "- Crear un dashboard simple con insights útiles para el equipo de marketing.", 
        "- Identificar perfiles con mayor probabilidad de recompra." 
    ], 
    "7. Preguntas para Discusión o Informe": [ 
        "- ¿Qué variables son más importantes para predecir la recompra?", 
        "- ¿Qué tipo de cliente conviene incentivar?", 
        "- ¿Cómo comunicarías tus hallazgos a alguien sin conocimientos técnicos?" 
    ], 
    "Entregables (opcional)": [ 
        "- Código en Jupyter Notebook o Python Script.", 
        "- Visualizaciones relevantes.", 
        "- Informe con introducción, análisis, resultados, y recomendaciones." 
    ], 
    "Rúbrica de Evaluación (opcional)": [ 
        "- Análisis Exploratorio: 20%", 
        "- Preprocesamiento de datos: 15%", 
        "- Modelado predictivo: 25%", 
        "- Visualización: 15%", 
        "- Interpretación y comunicación de resultados: 25%" 
    ] 
}

# Agregar secciones al documento 
for heading, content in sections.items(): 
    doc.add_heading(heading, level=1) 
    for paragraph in content: 
        doc.add_paragraph(paragraph)

# Guardar el documento
doc.save("Ejercicio_Mineria_Datos.docx")
